from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import _Cell

def create_template():
    doc = Document()
    
    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Republic of the Philippines\nDepartment of Social Welfare and Development\nField Office XI\n')
    run.font.size = Pt(10)
    run = p.add_run('HOME FOR THE AGED\n')
    run.bold = True
    run.font.size = Pt(12)
    run = p.add_run('Visayan Village, Tagum City, Davao del Norte')
    run.font.size = Pt(9)

    doc.add_paragraph() # Spacer

    # Title
    p = doc.add_paragraph('INVENTORY OF BELONGINGS')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)

    p = doc.add_paragraph('UPON ADMISSION')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    doc.add_paragraph()

    # Client Info
    p = doc.add_paragraph()
    p.add_run('Name of the client: ').bold = True
    p.add_run('{{ client_name }}').underline = True

    p = doc.add_paragraph()
    p.add_run('Date: ').bold = True
    p.add_run('{{ inventory_date }}').underline = True

    doc.add_paragraph()

    # Table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Headers
    headers = ['PARTICULARS', 'QTY', 'UNIT', 'DESCRIPTION', 'UNIT COST', 'BALANCE AS OF']
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row Loop
    # We add ONE row that contains the tags.
    # IMPORTANT: The {%p for %} and {%p endfor %} must be placed carefully.
    # docxtpl recommends placing the opening tag in the first cell and closing in the last cell of the row
    # BUT for row-wise loops, we usually put the tag *inside* the cell text or use specific methods.
    # The safest way for simple row repetition is:
    # Cell 1: {%p for item in admission_items %}{{ item.particulars }}
    # ...
    # Cell 6: {{ item.balance }}{%p endfor %}
    
    row = table.add_row()
    cells = row.cells
    
    cells[0].text = '{%p for item in admission_items %}{{ item.particulars }}'
    cells[1].text = '{{ item.qty }}'
    cells[2].text = '{{ item.unit }}'
    cells[3].text = '{{ item.description }}'
    cells[4].text = '{{ item.unit_cost }}'
    cells[5].text = '{{ item.balance }}{%p endfor %}'

    doc.add_paragraph()

    # Signatories
    p = doc.add_paragraph('Signature of Client: _______________________')
    doc.add_paragraph()

    # Signature Table (Implicit layout)
    sig_table = doc.add_table(rows=2, cols=4)
    sig_table.autofit = True
    # We'll just put placeholders
    curr_cells = sig_table.rows[0].cells
    curr_cells[0].text = 'Referring Party'
    curr_cells[1].text = 'Inspected by (HP)'
    curr_cells[2].text = 'Attested by (Supervising HP)'
    curr_cells[3].text = 'Noted by (Center Head)'

    val_cells = sig_table.rows[1].cells
    val_cells[0].text = '\n\n{{ referring_party }}'
    val_cells[1].text = '\n\n{{ inspected_by }}'
    val_cells[2].text = '\n\n{{ attested_by }}'
    val_cells[3].text = '\n\n{{ noted_by }}'
    
    # Save
    path = r"C:\eldercare_rcfms\form_templates\Home Life Service\inventory_admission.docx"
    doc.save(path)
    print(f"Created clean template at: {path}")

if __name__ == "__main__":
    create_template()
