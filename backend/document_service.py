from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Pt
import os
import io
import tempfile
import time
import requests
try:
    from docx2pdf import convert
    import pythoncom
    WINDOWS_PDF_AVAILABLE = True
except ImportError:
    convert = None
    pythoncom = None
    WINDOWS_PDF_AVAILABLE = False
import threading
from PIL import Image

class DocumentService:
    _thread_local = threading.local()

    @staticmethod
    def _get_word_app():
        if not WINDOWS_PDF_AVAILABLE:
            raise RuntimeError("Word COM automation not available on this platform")
        if not hasattr(DocumentService._thread_local, 'word_app'):
            pythoncom.CoInitialize()
            import win32com.client
            app = win32com.client.Dispatch("Word.Application")
            app.Visible = False
            app.DisplayAlerts = 0
            DocumentService._thread_local.word_app = app
        return DocumentService._thread_local.word_app

    @staticmethod
    def generate(template_type, service_unit, data, output_format='pdf'):
        """
        Generates a document by filling a Word template.
        Returns the path to the generated PDF or DOCX file.
        """
        try:
            # Initialize COM for docx2pdf (Windows only)
            if WINDOWS_PDF_AVAILABLE and pythoncom:
                pythoncom.CoInitialize()

            # 1. Resolve Template Path
            # Base dir is the backend folder's parent (root of project)
            current_dir = os.path.dirname(os.path.abspath(__file__))
            base_dir = os.path.abspath(os.path.join(current_dir, '..'))
            
            print(f" * DEBUG: Current Dir: {current_dir}")
            print(f" * DEBUG: Base Dir (Project Root): {base_dir}")

            template_path = DocumentService._find_template(base_dir, service_unit, template_type)
            
            if not template_path:
                # LIST DIRECTORIES FOR DEBUGGING
                templates_dir = os.path.join(base_dir, 'form_templates')
                if os.path.exists(templates_dir):
                     print(f" * DEBUG: contents of form_templates: {os.listdir(templates_dir)}")
                     target_dir = os.path.join(templates_dir, service_unit)
                     if os.path.exists(target_dir):
                         print(f" * DEBUG: contents of {service_unit}: {os.listdir(target_dir)}")
                     else:
                         print(f" * DEBUG: {service_unit} folder NOT FOUND in {templates_dir}")
                else:
                    print(f" * DEBUG: form_templates folder NOT FOUND at {templates_dir}")

                raise FileNotFoundError(f"Template not found for type: {template_type} in {service_unit}")

            # SPECIAL INJECTION FOR OUT ON PASS DYNAMIC DESIGNATIONS
            if template_type in ['out_on_pass', 'homelife_out_on_pass']:
                # Import supabase here to avoid circular imports if needed, or use global
                try:
                    from app import supabase
                    # Center Doctor Designation
                    cd_name = data.get('center_doctor')
                    if cd_name and not data.get('center_doctor_designation'):
                        profile = supabase.table('profiles').select('job_title, role').ilike('full_name', cd_name).execute()
                        if profile.data:
                            jt = profile.data[0].get('job_title')
                            role = profile.data[0].get('role')
                            if not jt:
                                # Fallback to nice formatting of role if no job title
                                jt = role.replace('_', ' ').title() if role else 'Medical Officer'
                            data['center_doctor_designation'] = str(jt).upper()
                    
                    # Social Worker Designation
                    sw_name = data.get('social_worker')
                    if sw_name and not data.get('social_worker_designation'):
                        profile = supabase.table('profiles').select('job_title, role').ilike('full_name', sw_name).execute()
                        if profile.data:
                            jt = profile.data[0].get('job_title')
                            role = profile.data[0].get('role')
                            if not jt:
                                jt = role.replace('_', ' ').title() if role else 'Social Worker'
                            data['social_worker_designation'] = str(jt).upper()
                except Exception as e:
                    print(f" * WARNING: Failed to dynamically fetch designations for Out on Pass: {e}")

            print(f" * Loading template: {template_path}")
            t_start = time.time()
            doc = DocxTemplate(template_path)

            # 2. Pre-process data (format currency, resolve signatures)
            context = DocumentService._format_data(data)
            DocumentService._resolve_signature_images(doc, context)
            t_load = time.time()
            print(f" * [PERF] Load & Pre-process: {t_load - t_start:.2f}s")

            # 3. Render Template
            doc.render(context)
            t_render = time.time()
            print(f" * [PERF] doc.render(): {t_render - t_load:.2f}s")

            # 3a. Zero out paragraph spacing around signature images
            DocumentService._fix_signature_paragraph_spacing(doc)

            # 4. Optimization: Prune trailing empty paragraphs to avoid extra PDF pages
            DocumentService._remove_trailing_empty_paragraphs(doc)

            # 5. Save to Temporary DOCX File
            timestamp = int(time.time())
            docx_filename = f"generated_{template_type}_{timestamp}.docx"
            temp_dir = tempfile.gettempdir()
            docx_path = os.path.join(temp_dir, docx_filename)
            
            doc.save(docx_path)
            print(f" * DOCX saved to: {docx_path}")
            t_save = time.time()
            print(f" * [PERF] docx processing & save: {t_save - t_render:.2f}s")

            # 6. Return DOCX directly if requested
            if output_format == 'docx':
                print(f" * Returning DOCX (output_format=docx)")
                return docx_path

            # 7. Convert to PDF using Cached COM
            pdf_filename = f"generated_{template_type}_{timestamp}.pdf"
            pdf_path = os.path.join(temp_dir, pdf_filename)
            
            print(f" * Converting to PDF: {pdf_path}")
            try:
                # Use cached COM for much faster conversion
                t_com_start = time.time()
                word = DocumentService._get_word_app()
                abs_docx_path = os.path.abspath(docx_path)
                abs_pdf_path = os.path.abspath(pdf_path)
                
                doc_com = word.Documents.Open(abs_docx_path)
                t_com_open = time.time()
                
                doc_com.SaveAs(abs_pdf_path, FileFormat=17) # 17 = wdFormatPDF
                t_com_save = time.time()
                
                doc_com.Close(0) # 0 = wdDoNotSaveChanges
                t_com_close = time.time()
                
                print(f" * [PERF] COM Open: {t_com_open - t_com_start:.2f}s, SaveAs PDF: {t_com_save - t_com_open:.2f}s, Close: {t_com_close - t_com_save:.2f}s")
                print(f" * [PERF] TOTAL COM TIME: {t_com_close - t_com_start:.2f}s")
                
                if os.path.exists(pdf_path):
                    print(f" * PDF Conversion Successful (Fast Cached COM)")
                    return pdf_path
                else:
                    print(" * PDF Conversion failed (File not created)")
                    return docx_path # Fallback to DOCX
                    
            except Exception as e:
                print(f" * Fast PDF Conversion Error: {e}")
                print(" * Falling back to docx2pdf...")
                try:
                    t_fall_start = time.time()
                    from docx2pdf import convert
                    convert(docx_path, pdf_path)
                    print(f" * [PERF] Fallback docx2pdf time: {time.time() - t_fall_start:.2f}s")
                    if os.path.exists(pdf_path):
                        return pdf_path
                except Exception as e2:
                    print(f" * Fallback Conversion Error: {e2}")
                return docx_path

        except Exception as e:
            print(f"Error in DocumentService: {e}")
            raise e

    @staticmethod
    def _fix_signature_paragraph_spacing(doc):
        """
        After rendering, finds every paragraph that contains an inline image
        (i.e. a signature) and zeroes out its Space Before, Space After, and
        line spacing so the signature sits flush against adjacent text.
        Additionally, shifts the image downward slightly to overlap the text below it.
        """
        try:
            from docx.oxml import parse_xml
            for para in doc.paragraphs:
                # Check if any run's XML contains an embedded image blip
                has_image = any(
                    '<a:blip' in run._element.xml
                    for run in para.runs
                    if run._element is not None
                )
                if has_image:
                    fmt = para.paragraph_format
                    fmt.space_before = Pt(0)
                    fmt.space_after = Pt(0)
                    fmt.line_spacing = Pt(1)
                    
                    # Apply physical signature overlap effect
                    for run in para.runs:
                        if '<a:blip' in run._element.xml:
                            rPr = run._element.get_or_add_rPr()
                            # Shift the run DOWN by 60 points (-120 half-points)
                            # This creates a deeper overlap over the text below it
                            position = parse_xml('<w:position w:val="-120" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                            rPr.append(position)
                            
                    print(f" * Zeroed spacing and applied physical overlap on signature")
        except Exception as e:
            print(f" * WARNING: Failed to fix signature paragraph spacing: {e}")

    @staticmethod
    def _crop_transparent_signature(image_bytes):
        """
        Crops all the empty transparent or white padding around a drawn signature PNG
        so it doesn't create huge gaps when inserted into a Word document.
        """
        try:
            img = Image.open(io.BytesIO(image_bytes)).convert("RGBA")
            
            # Create a white background image
            bg = Image.new("RGBA", img.size, (255, 255, 255, 255))
            # Paste the signature on top of the white background
            # This turns transparent areas into white
            bg.paste(img, (0, 0), img)
            
            # Convert to grayscale to evaluate lightness
            gray = bg.convert("L")
            
            # Threshold: anything darker than near-white (250) is part of the signature
            # This mask will be 255 (white/solid) for signature pixels and 0 for background
            mask = gray.point(lambda p: 255 if p < 250 else 0)
            
            # Get the bounding box of the non-zero (signature) regions
            bbox = mask.getbbox()
            
            if bbox:
                # Crop the original image (preserving any intentional transparency within the bbox)
                img = img.crop(bbox)
            
            out = io.BytesIO()
            img.save(out, format="PNG")
            return out.getvalue()
        except Exception as e:
            print(f" * WARNING: Failed to crop signature: {e}")
            return image_bytes

    @staticmethod
    def _format_data(data):
        """
        Recursively formats fields in the data dictionary.
        Specifically handles currency fields (unit_cost, balance) to ensure .00 formatting.
        """
        if isinstance(data, list):
            return [DocumentService._format_data(item) for item in data]
        if isinstance(data, dict):
            new_data = {}
            for key, value in data.items():
                # Fields to format as currency if they are numeric
                if key in ['unit_cost', 'balance'] and value is not None:
                    try:
                        # Only format if it's numeric
                        num_val = float(value)
                        new_data[key] = f"{num_val:.2f}"
                    except (ValueError, TypeError):
                        new_data[key] = DocumentService._format_data(value)
                else:
                    new_data[key] = DocumentService._format_data(value)
            return new_data
        return data

    @staticmethod
    def _resolve_signature_images(doc, data):
        """
        Scans the data dict for keys ending in '_signature_url'.
        Downloads each image, crops empty padding, and replaces the URL string 
        with a docxtpl InlineImage object.
        """
        if not data or not isinstance(data, dict):
            return

        signature_keys = [k for k in data.keys() if k.endswith('_signature_url')]

        # Proactive mapping: if we have 'center_head_name_signature_url', also provide 'center_head_signature_url'
        # and vice versa, as different templates use different conventions.
        if 'center_head_name_signature_url' in data and 'center_head_signature_url' not in data:
            data['center_head_signature_url'] = data['center_head_name_signature_url']
        if 'center_head_signature_url' in data and 'center_head_name_signature_url' not in data:
            data['center_head_name_signature_url'] = data['center_head_signature_url']
        
        # New: Harmonize 'noted_by_signature_url' with Center Head aliases
        if 'noted_by_signature_url' in data:
            if 'center_head_signature_url' not in data:
                data['center_head_signature_url'] = data['noted_by_signature_url']
            if 'center_head_name_signature_url' not in data:
                data['center_head_name_signature_url'] = data['noted_by_signature_url']
        elif 'center_head_signature_url' in data:
            data['noted_by_signature_url'] = data['center_head_signature_url']
        
        # Refresh keys after mapping
        signature_keys = [k for k in data.keys() if k.endswith('_signature_url')]

        for key in signature_keys:
            url = data.get(key)
            if not url or not isinstance(url, str):
                data[key] = ''
                continue

            try:
                content = None
                if url.startswith('http'):
                    print(f" * Downloading and cropping signature: {key} -> {url[:80]}...")
                    resp = requests.get(url, timeout=15)
                    if resp.status_code == 200 and len(resp.content) > 0:
                        content = resp.content
                    else:
                        print(f" * WARNING: Failed to download {key}: HTTP {resp.status_code}")
                elif url.startswith('data:image'):
                    print(f" * Decoding base64 signature for {key}...")
                    try:
                        import base64
                        header, encoded = url.split(",", 1)
                        content = base64.b64decode(encoded)
                    except Exception as b64e:
                        print(f" * WARNING: Failed to decode base64 for {key}: {b64e}")
                else:
                    print(f" * WARNING: Invalid signature URL/URI for {key}")

                if content:
                    cropped_bytes = DocumentService._crop_transparent_signature(content)
                    image_stream = io.BytesIO(cropped_bytes)
                    
                    # Calculate aspect ratio to dynamically set width/height
                    with Image.open(image_stream) as img:
                        w, h = img.size
                    # Reset stream position after reading
                    image_stream.seek(0)
                    
                    # Maximum allowed dimensions (in mm)
                    MAX_WIDTH = 40.0
                    MAX_HEIGHT = 16.0
                    
                    aspect_ratio = w / h if h > 0 else 1.0
                    
                    # If the image is very wide relative to its height (like a typical drawn signature)
                    if aspect_ratio >= (MAX_WIDTH / MAX_HEIGHT):
                        # Bound by width
                        data[key] = InlineImage(doc, image_stream, width=Mm(MAX_WIDTH))
                    else:
                        # Bound by height (like a tall uploaded image)
                        data[key] = InlineImage(doc, image_stream, height=Mm(MAX_HEIGHT))
                        
                    print(f" * Signature cropped & injected: {key} (Aspect Ratio: {aspect_ratio:.2f})")
                else:
                    data[key] = ''

            except Exception as e:
                print(f" * WARNING: Error resolving signature {key}: {e}")
                data[key] = ''  # Graceful fallback

    @staticmethod
    def _remove_trailing_empty_paragraphs(doc):
        """
        Removes empty trailing paragraphs from the end of the document.
        This prevents unnecessary blank pages when converting to PDF.
        """
        try:
            # We iterate backwards through paragraphs and remove those that are purely whitespace
            # and have no sub-elements like images or tables.
            while doc.paragraphs:
                last_p = doc.paragraphs[-1]
                # Check for text and sub-elements (like runs with images)
                if not last_p.text.strip() and not last_p.runs:
                    p = last_p._element
                    p.getparent().remove(p)
                else:
                    break
        except Exception as e:
            # Failure here should not block the overall generation
            print(f" * WARNING: Failed to prune trailing paragraphs: {e}")

    @staticmethod
    def _find_template(base_dir, service_unit, template_type):
        """
        Helper to find the correct .docx file.
        service_unit: 'Social Service', 'Home Life Service', etc.
        template_type: 'admission_slip', 'case_conference', etc.
        """
        folder_name = service_unit
        
        # 1. Try exact match in known folder
        candidate_path = os.path.join(base_dir, 'form_templates', folder_name, f"{template_type}.docx")
        if os.path.exists(candidate_path):
            return candidate_path

        # 2. Fuzzy search in specific folder
        target_folder = os.path.join(base_dir, 'form_templates', folder_name)
        if os.path.exists(target_folder):
            print(f" * DEBUG: Searching in {target_folder} for {template_type}")
            for file in os.listdir(target_folder):
                if file.startswith('~'): continue # Ignore temp files
                
                # Check 1: Starts with template_type (ignoring case)
                if file.lower().startswith(template_type.lower()):
                     print(f" * Found match: {file}")
                     return os.path.join(target_folder, file)
                
                # Check 2: Starts with template_type (replacing underscores with spaces)
                clean_type = template_type.lower().replace('_', ' ')
                if file.lower().startswith(clean_type):
                     print(f" * Found match (clean): {file}")
                     return os.path.join(target_folder, file)

        return None
